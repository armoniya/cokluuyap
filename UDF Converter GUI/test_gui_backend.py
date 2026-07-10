import os
import sys
import zipfile
import datetime
from cryptography import x509
from cryptography.x509.oid import NameOID
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.asymmetric import rsa, padding
from cryptography.hazmat.primitives.serialization import Encoding

sys.path.append(os.path.dirname(os.path.abspath(__file__)))

import converter
import cades
import verify

def create_mock_docx(file_path):
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
    
    doc = docx.Document()
    
    # 1. Title Paragraph (Centered, Bold, Size 16)
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(12)
    run1 = p1.add_run("T.C. ANKARA NÖBETÇİ ASLİYE HUKUK MAHKEMESİNE")
    run1.bold = True
    run1.font.name = "Times New Roman"
    run1.font.size = Pt(14)
    
    # 2. Case Info Paragraph (Right Aligned, Bold, Size 11)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_after = Pt(18)
    run2 = p2.add_run("DOSYA NO: 2026/123 Esas")
    run2.bold = True
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(11)
    
    # 3. Main Body Paragraph (Justified, Custom Line Spacing, First Line Indent)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p3.paragraph_format.first_line_indent = Pt(35.4) # ~1.25 cm
    p3.paragraph_format.line_spacing = 1.5
    p3.paragraph_format.space_after = Pt(6)
    
    run3_1 = p3.add_run("Davacı vekili olarak, mahkemenizin yukarıda belirtilen esasına kayden görülmekte olan dava dosyası kapsamında ")
    run3_1.font.name = "Times New Roman"
    run3_1.font.size = Pt(12)
    
    run3_2 = p3.add_run("delil listemizi ")
    run3_2.bold = True
    run3_2.font.name = "Times New Roman"
    run3_2.font.size = Pt(12)
    
    run3_3 = p3.add_run("sunmaktan şeref duyarız. Ekli belgelerin ve tanık listesinin dosyaya celbini talep ederiz. ")
    run3_3.italic = True
    run3_3.font.name = "Times New Roman"
    run3_3.font.size = Pt(12)
    
    # 4. Table to test table extraction
    table = doc.add_table(rows=2, cols=2)
    cell_1 = table.cell(0, 0)
    cell_1.text = "Sıra No"
    cell_2 = table.cell(0, 1)
    cell_2.text = "Açıklama"
    cell_3 = table.cell(1, 0)
    cell_3.text = "1"
    cell_4 = table.cell(1, 1)
    cell_4.text = "Banka Dekontu (Ek-1)"
    
    doc.save(file_path)
    print(f"-> Test amaçlı DOCX oluşturuldu: {file_path}")

def create_mock_certificate_and_key():
    private_key = rsa.generate_private_key(
        public_exponent=65537,
        key_size=2048,
    )
    
    subject = issuer = x509.Name([
        x509.NameAttribute(NameOID.COUNTRY_NAME, "TR"),
        x509.NameAttribute(NameOID.STATE_OR_PROVINCE_NAME, "Ankara"),
        x509.NameAttribute(NameOID.ORGANIZATION_NAME, "UDF Converter Test Kurumu"),
        x509.NameAttribute(NameOID.COMMON_NAME, "Mock Avukat"),
    ])
    
    now = datetime.datetime.now(datetime.timezone.utc)
    cert = x509.CertificateBuilder().subject_name(
        subject
    ).issuer_name(
        issuer
    ).public_key(
        private_key.public_key()
    ).serial_number(
        x509.random_serial_number()
    ).not_valid_before(
        now - datetime.timedelta(days=1)
    ).not_valid_after(
        now + datetime.timedelta(days=365)
    ).add_extension(
        x509.KeyUsage(
            digital_signature=True,
            content_commitment=True,
            key_encipherment=False,
            data_encipherment=False,
            key_agreement=False,
            key_cert_sign=False,
            crl_sign=False,
            encipher_only=False,
            decipher_only=False
        ),
        critical=True
    ).sign(private_key, hashes.SHA256())
    
    cert_der = cert.public_bytes(Encoding.DER)
    return private_key, cert_der

def test_conversion_and_mock_sign():
    docx_path = "test_document.docx"
    udf_path = "test_document_mock_signed.udf"
    
    # 1. Create mock DOCX
    create_mock_docx(docx_path)
    
    # 2. Convert to UDF XML
    print("2) DOCX'ten UDF XML dönüştürülüyor...")
    xml_bytes = converter.docx_to_udf_xml(docx_path)
    print(f"UDF XML boyutu: {len(xml_bytes)} bayt")
    
    # 3. Create Mock Certificate
    priv_key, cert_der = create_mock_certificate_and_key()
    
    # Build CMS envelope (sign.sgn)
    print("3) CAdES-BES yapısı ve mock e-imza oluşturuluyor...")
    signed_attrs, _, _ = cades.to_be_signed(xml_bytes, cert_der)
    
    # Sign signed_attrs directly - cryptography's sign() with hashes.SHA256()
    # will automatically hash it and apply the DigestInfo prefix and PKCS1v15 padding.
    signature = priv_key.sign(
        signed_attrs,
        padding.PKCS1v15(),
        hashes.SHA256()
    )
    
    # Build CMS envelope (sign.sgn)
    sgn = cades.build_cms(xml_bytes, cert_der, signed_attrs, signature, use_ecdsa=False)
    
    # Package into UDF
    with zipfile.ZipFile(udf_path, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("content.xml", xml_bytes)
        z.writestr("sign.sgn", sgn)
        
    print(f"-> Mock imzalanmış UDF paketi oluşturuldu: {udf_path}")
    
    # 5. Verify the package
    print("\n4) Üretilen UDF paketi bütünlük ve imza testine sokuluyor...")
    success, logs = verify.verify_udf_log(udf_path)
    
    print("\n=== Doğrulama Günlüğü ===")
    for log in logs:
        print(log)
        
    if success:
        print("\n[BAŞARILI] UDF paketi doğrulamadan sorunsuz geçti!")
    else:
        print("\n[HATA] UDF paketi doğrulamadan geçemedi.")
        
    # Clean temporary files
    try:
        os.remove(docx_path)
        os.remove(udf_path)
        print("Temizlik tamamlandı.")
    except Exception:
        pass

if __name__ == "__main__":
    test_conversion_and_mock_sign()
